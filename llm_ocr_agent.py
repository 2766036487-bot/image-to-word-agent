import os
import uuid
import json
from dotenv import load_dotenv
from fastapi import FastAPI, UploadFile, File, HTTPException, Form
from fastapi.responses import FileResponse
import easyocr
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from openai import OpenAI

# ========== 1. 加载环境变量（核心：保护敏感信息） ==========
load_dotenv()  # 读取同目录的.env文件

# ========== 2. 基础配置（统一从环境变量读取，无硬编码） ==========
# EasyOCR模型路径
os.environ['EASYOCR_MODULE_PATH'] = os.getenv("EASYOCR_MODEL_PATH")
# 临时文件目录（从.env读取，适配不同环境）
UPLOAD_DIR = os.getenv("UPLOAD_DIR")
OUTPUT_DIR = os.getenv("OUTPUT_DIR")

# 确保目录存在（兼容不同系统路径）
os.makedirs(UPLOAD_DIR, exist_ok=True)
os.makedirs(OUTPUT_DIR, exist_ok=True)

# ========== 3. 初始化核心组件（无硬编码） ==========
# 初始化EasyOCR工具
ocr_reader = easyocr.Reader(['ch_sim', 'en'], gpu=False)

# 初始化智谱GLM-4.5-Flash客户端（从.env读API Key）
llm_client = OpenAI(
    api_key=os.getenv("GLM_API_KEY"),  # 无硬编码！
    base_url="https://open.bigmodel.cn/api/paas/v4"
)

# 初始化FastAPI应用
app = FastAPI(
    title="GLM-4.5-Flash驱动的图片转Word智能体",
    description="学生专用：大模型自主调度OCR工具生成Word（免费版）"
)


# ========== 4. 工具函数：图片转Word（复用验证过的逻辑） ==========
def image_to_word_tool(image_path: str, output_path: str) -> bool:
    """
    图片转Word工具函数（供大模型调用）
    :param image_path: 图片绝对路径
    :param output_path: Word输出绝对路径
    :return: 成功返回True，失败返回False
    """
    try:
        # 1. OCR识别图片文字
        result = ocr_reader.readtext(image_path, detail=1, paragraph=False)

        # 2. 按坐标排序（还原图片文字顺序）
        def sort_key(item):
            return (item[0][0][1], item[0][0][0])  # y坐标优先，再x坐标

        result_sorted = sorted(result, key=sort_key)

        # 3. 合并同行文字
        merged_lines, current_line, current_y = [], "", None
        y_threshold = 15  # 同一行y坐标误差阈值
        for item in result_sorted:
            if len(item) >= 3 and item[2] > 0.4:  # 过滤低置信度结果
                text, y1 = item[1].strip(), item[0][0][1]
                if current_y is None or abs(y1 - current_y) <= y_threshold:
                    current_line += text + " "
                else:
                    if current_line.strip():
                        merged_lines.append(current_line.strip())
                    current_line = text + " "
                current_y = y1
        if current_line.strip():
            merged_lines.append(current_line.strip())

        # 4. 区分普通文字和表格文字
        normal_text, table_text, seen_lines = [], [], set()
        for line in merged_lines:
            if line not in seen_lines:
                seen_lines.add(line)
                cols = [c.strip() for c in line.split("|") if c.strip()]
                if "|" in line and len(cols) >= 2:
                    table_text.append(line)
                else:
                    normal_text.append(line)

        # 5. 生成Word文档（统一设置为宋体）
        doc = Document()
        # 设置全局宋体样式
        style = doc.styles['Normal']
        style.font.name = '宋体'
        style.font.size = Pt(12)
        style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        # 写入普通文字
        for line in normal_text:
            para = doc.add_paragraph(line)
            para.runs[0].font.name = '宋体'
            para.runs[0].element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        # 写入表格
        if table_text:
            doc.add_paragraph()  # 空行分隔
            table_data = [[c.strip() for c in line.split("|") if c.strip()] for line in table_text]
            table_data = [cols for cols in table_data if cols]
            if table_data:
                rows, cols = len(table_data), max(len(row) for row in table_data)
                table = doc.add_table(rows=rows, cols=cols)
                table.style = 'Table Grid'  # 带边框表格
                for i in range(rows):
                    for j in range(len(table_data[i])):
                        if j < cols:
                            cell = table.rows[i].cells[j]
                            cell.text = table_data[i][j]
                            # 表格文字设为宋体
                            cell.paragraphs[0].runs[0].font.name = '宋体'
                            cell.paragraphs[0].runs[0].element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        # 保存Word文档
        doc.save(output_path)
        return True
    except Exception as e:
        print(f"图片转Word工具执行失败：{str(e)}")
        return False


# ========== 5. 大模型工具调用逻辑（GLM-4.5-Flash核心） ==========
def llm_tool_caller(user_query: str, image_filename: str = None) -> dict:
    """
    GLM-4.5-Flash判断是否调用工具，并执行对应逻辑
    :param user_query: 用户指令
    :param image_filename: 上传的图片文件名（None则无图片）
    :return: 处理结果字典
    """
    # 定义工具描述（告诉大模型可用工具）
    tools = [
        {
            "type": "function",
            "function": {
                "name": "image_to_word_tool",
                "description": "将图片中的文字和表格识别并转换为Word文档，仅当用户要求处理图片时调用",
                "parameters": {
                    "type": "object",
                    "properties": {
                        "image_path": {"type": "string", "description": "图片绝对路径"},
                        "output_path": {"type": "string", "description": "Word输出绝对路径"}
                    },
                    "required": ["image_path", "output_path"]
                }
            }
        }
    ]

    # 大模型系统指令（强化学生场景适配）
    system_prompt = f"""
    你是一个运行在学生本地环境的图片转Word智能体，核心规则：
    1. 仅当用户上传图片且要求转Word时，调用image_to_word_tool工具；
    2. 自动填充参数：
       - image_path = {UPLOAD_DIR}/{image_filename}
       - output_path = {OUTPUT_DIR}/{uuid.uuid4()}_result.docx
    3. 无需询问用户参数，直接调用工具；
    4. 回复简洁，仅返回核心结果（下载链接/失败提示）。
    """

    # 调用GLM-4.5-Flash模型
    try:
        response = llm_client.chat.completions.create(
            model="glm-4.5-flash",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user",
                 "content": f"用户指令：{user_query}，已上传图片：{image_filename if image_filename else '无'}"}
            ],
            tools=tools,
            tool_choice="auto",
            extra_body={"thinking.type": "enabled"}  # 开启动态思考
        )
    except Exception as e:
        return {"code": 500, "message": f"大模型调用失败：{str(e)}"}

    # 解析模型响应
    response_message = response.choices[0].message
    tool_calls = response_message.tool_calls

    # 情况1：模型决定调用工具
    if tool_calls and image_filename:
        file_uuid = str(uuid.uuid4())
        image_path = os.path.join(UPLOAD_DIR, image_filename)
        word_filename = f"{file_uuid}_result.docx"
        word_path = os.path.join(OUTPUT_DIR, word_filename)

        # 执行图片转Word工具
        tool_success = image_to_word_tool(image_path, word_path)
        if tool_success:
            return {
                "code": 200,
                "message": "✅ GLM-4.5-Flash已成功生成Word文档！",
                "word_filename": word_filename,
                "download_url": f"http://127.0.0.1:8000/download-word/?filename={word_filename}"
            }
        else:
            return {"code": 500, "message": "❌ 图片识别失败，请检查图片是否清晰！"}

    # 情况2：无图片/模型不调用工具
    else:
        return {
            "code": 200,
            "message": "💡 GLM-4.5-Flash回复：",
            "llm_response": response_message.content or "请上传图片并输入“图片转Word”相关指令！"
        }


# ========== 6. Web接口（用户交互层） ==========
@app.post("/agent/process/", summary="上传图片+指令，生成Word")
async def agent_process(
        user_query: str = Form(description="用户指令，如'把这张图片转成Word'"),
        file: UploadFile = File(None, description="需要处理的图片（jpg/png/jpeg）")
):
    """智能体核心交互接口"""
    try:
        image_filename = None
        # 保存上传的图片
        if file:
            allowed_types = ["image/jpeg", "image/png", "image/jpg"]
            if file.content_type not in allowed_types:
                raise HTTPException(status_code=400, detail="仅支持jpg/png/jpeg格式的图片！")
            # 生成唯一文件名（避免重复）
            image_filename = f"{uuid.uuid4()}_{file.filename}"
            image_path = os.path.join(UPLOAD_DIR, image_filename)
            with open(image_path, "wb") as f:
                f.write(await file.read())

        # 调用大模型工具调度逻辑
        result = llm_tool_caller(user_query, image_filename)
        return result

    except HTTPException as e:
        raise e
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"智能体处理失败：{str(e)}")


@app.get("/download-word/", summary="下载生成的Word文档")
async def download_word(filename: str):
    """下载Word文档接口"""
    try:
        word_path = os.path.join(OUTPUT_DIR, filename)
        if not os.path.exists(word_path):
            raise HTTPException(status_code=404, detail="Word文件不存在！")
        return FileResponse(
            path=word_path,
            filename=filename,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    except HTTPException as e:
        raise e
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"下载失败：{str(e)}")


# ========== 7. 启动服务 ==========
if __name__ == "__main__":
    import uvicorn

    # 启动FastAPI服务（本地访问：http://127.0.0.1:8000）
    #http://127.0.0.1:8000/docs
    uvicorn.run(app, host="0.0.0.0", port=8000)